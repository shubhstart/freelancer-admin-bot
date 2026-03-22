
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_slide_header(doc, title, subtitle=None):
    # Add a page break for the next "slide"
    doc.add_page_break()
    
    # Title
    t = doc.add_heading(title, level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x3B) # Dark Blue-Gray
        run.font.name = 'Arial'

    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in sub.runs:
            run.font.size = Pt(14)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B) # Muted Gray

    doc.add_paragraph("\n")

def create_slides():
    doc = Document()

    # --- SLIDE 1: TITLE ---
    title_p = doc.add_paragraph("\n\n\n\n")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("FREELANCER ADMIN AUTOMATION")
    run.font.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Dark Navy

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("Implementation Overview & Multi-Agent Architecture")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    doc.add_paragraph("\n\n")
    company_p = doc.add_paragraph()
    company_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company_p.add_run("By Hallucination Hunters")
    run.font.size = Pt(14)
    run.font.bold = True

    # --- SLIDE 2: ARCHITECTURE ---
    add_slide_header(doc, "System Architecture", "The Orchestrator-Agent Pattern")
    
    doc.add_paragraph("The application follows a Hub-and-Spoke model where a central Flask orchestrator routes user requests to specialized task-specific agents.")
    
    # Textual Diagram
    diagram = doc.add_paragraph()
    diagram.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = diagram.add_run(
        "┌───────────────────┐\n"
        "│  User Interface   │\n"
        "└─────────┬─────────┘\n"
        "          ▼\n"
        "┌───────────────────┐\n"
        "│ Flask Orchestrator│ (app.py)\n"
        "└────┬─────┬─────┬──┘\n"
        "     ▼     ▼     ▼\n"
        " ┌──────┐┌──────┐┌──────┐\n"
        " │Props ││Invs  ││Remind│ (Specialized Agents)\n"
        " └──────┘└──────┘└──────┘"
    )
    run.font.name = 'Courier New'
    run.font.size = Pt(11)

    # --- SLIDE 3: INTENT CLASSIFICATION ---
    add_slide_header(doc, "Intent Classification", "AI-Driven Routing")
    
    doc.add_paragraph("Using LLM-based zero-shot classification to detect what the freelancer wants to do.")
    
    p = doc.add_paragraph()
    p.add_run("Core Logic: ").font.bold = True
    p.add_run("The `intent.py` module uses GPT-4o-mini (or Llama 3.2) with a strict system prompt to return one of five labels: ")
    p.add_run("PROPOSAL, INVOICE, REMINDER, QUERY, or GENERAL.").font.italic = True

    code = doc.add_paragraph()
    run = code.add_run(
        "INTENT_SYSTEM = \"\"\"Reply with EXACTLY ONE of these labels:\n"
        "PROPOSAL, INVOICE, REMINDER, QUERY, GENERAL\n"
        "CRITICAL: Return ONLY THE LABEL.\"\"\""
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

    # --- SLIDE 4: PROPOSAL AGENT ---
    add_slide_header(doc, "Proposal Generator", "Automated Document Creation")
    
    doc.add_paragraph("Converts a project description into a multi-section professional proposal.")
    
    bullets = ["Extracts: Client Name, Project Title, Description, Deliverables, Timeline",
               "Generates: Professional introduction, terms, pricing, and signature",
               "Exports: Both PDF (ReportLab) and DOCX (python-docx)"]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # --- SLIDE 5: INVOICE AGENT ---
    add_slide_header(doc, "Invoice Generator", "Calculation & PDF Export")
    
    doc.add_paragraph("Handles billing math and professional formatting.")
    
    doc.add_paragraph("Key Workflow:", style='Heading 2')
    doc.add_paragraph("1. Collect hours and rates from user chat.", style='List Number')
    doc.add_paragraph("2. Calculate Subtotal, Tax (GST/VAT), and Grand Total.", style='List Number')
    doc.add_paragraph("3. Generate a professional table-based PDF.", style='List Number')
    doc.add_paragraph("4. Persist invoice status to SQLite for tracking.", style='List Number')

    # --- SLIDE 6: REMINDER AGENT ---
    add_slide_header(doc, "Payment Reminders", "Tone-Aware Automation")
    
    doc.add_paragraph("Detects how overdue an invoice is and adapts the communication tone accordingly.")
    
    doc.add_paragraph("Tone Tiers:", style='Heading 2')
    doc.add_paragraph("• 1-7 Days: Gentle (Friendly reminder)", style='Body Text')
    doc.add_paragraph("• 8-21 Days: Firm (Direct request)", style='Body Text')
    doc.add_paragraph("• 22+ Days: Urgent (Professional but stern)", style='Body Text')

    # --- SLIDE 7: FREECELERATOR (FREE VERSION) ---
    add_slide_header(doc, "Local LLM Integration", "Privacy & Cost-Free Execution")
    
    doc.add_paragraph("Modified for the 'Free' version to use local models via Ollama.")
    
    doc.add_paragraph("Configuration Details:", style='Heading 2')
    doc.add_paragraph("• Default Model: Llama 3.2", style='Body Text')
    doc.add_paragraph("• Backend: OpenAI-compatible API at http://localhost:11434/v1", style='Body Text')
    doc.add_paragraph("• Privacy: No data leaves the freelancer’s machine.", style='Body Text')

    # --- SLIDE 8: DATA PERSISTENCE ---
    add_slide_header(doc, "Database Layer", "SQLite & CRUD Operations")
    
    doc.add_paragraph("A relational database ensures data persistence across sessions.")
    
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Table Name'
    hdr_cells[1].text = 'Stored Entities'
    
    data = [('clients', 'Name, Email, Contact Info'),
            ('proposals', 'Text, File Paths, Timestamps'),
            ('invoices', 'Number, Totals, Status (UNPAID/PAID)'),
            ('reminders', 'Drafts and Sent Logs')]
    
    for tn, se in data:
        row_cells = table.add_row().cells
        row_cells[0].text = tn
        row_cells[1].text = se

    doc.save('Implementation_Overview_Slides.docx')
    print("Document created successfully.")

if __name__ == "__main__":
    create_slides()
