"""Proposal conversation agent — collects fields, generates doc via LLM, saves PDF + DOCX."""

import os
import json
from datetime import datetime
from openai import OpenAI
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import cm
from .. import database as db


REQUIRED_FIELDS = [
    "client_name", "project_title", "project_description",
    "deliverables", "timeline", "freelancer_name",
]
OPTIONAL_FIELDS = ["budget", "freelancer_skills"]
ALL_FIELDS = REQUIRED_FIELDS + OPTIONAL_FIELDS

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "generated_docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)

EXTRACT_SYSTEM = """You are a data-extraction assistant.
Given the conversation so far about a freelance proposal, extract any of these fields
that have been provided. Return ONLY a JSON object with the keys you can fill:

  client_name, project_title, project_description, deliverables,
  timeline, freelancer_name, budget, freelancer_skills

If a field has not been mentioned, omit it.  Reply with valid JSON only."""

GENERATE_SYSTEM = """You are a professional proposal writer for the company "Hallucination Hunters".
Given the structured data below, write a polished, client-ready project proposal
with these sections:
1. Introduction
2. Project Understanding
3. Proposed Approach
4. Deliverables
5. Timeline
6. Pricing
7. Terms
8. Closing (include company contact: hallucination_hunters@gmail.com | 08043672418 | hallucination_hunters.com)

Write in a professional, confident tone. Use markdown formatting."""


def _extract_fields(client: OpenAI, conversation: list, model="gpt-4o-mini") -> dict:
    msgs = [{"role": "system", "content": EXTRACT_SYSTEM}]
    # Condense conversation into one user message
    text = "\n".join(f"{m['role']}: {m['content']}" for m in conversation)
    msgs.append({"role": "user", "content": text})
    resp = client.chat.completions.create(model=model, temperature=0, max_tokens=600, messages=msgs)
    raw = resp.choices[0].message.content.strip()
    # Strip markdown code fences if present
    if raw.startswith("```"):
        raw = "\n".join(raw.split("\n")[1:])
    if raw.endswith("```"):
        raw = "\n".join(raw.split("\n")[:-1])
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {}


def _missing(fields: dict) -> list:
    return [f for f in REQUIRED_FIELDS if not fields.get(f)]


def _ask_next(missing: list) -> str:
    # Always show the full template so the user can fill everything in one shot
    template = (
        "Client Name: \n\n"
        "Project Title: \n\n"
        "Project Description: \n\n"
        "Deliverables: \n\n"
        "Timeline: \n\n"
        "Your Name (Freelancer): Hallucination Hunters\n\n"
        "Budget: (optional)\n\n"
        "Your Skills: AI, Full Stack Development"
    )
    return (
        "Please fill in the details below and send it back:\n\n"
        f"{template}"
    )


def _generate_proposal_text(client: OpenAI, fields: dict, model="gpt-4o-mini", lang="English") -> str:
    lang_note = f"\nIMPORTANT: Write the entire proposal in {lang} language." if lang != "English" else ""
    msgs = [
        {"role": "system", "content": GENERATE_SYSTEM + lang_note},
        {"role": "user", "content": json.dumps(fields, indent=2)},
    ]
    resp = client.chat.completions.create(model=model, temperature=0.7, max_tokens=2000, messages=msgs)
    return resp.choices[0].message.content.strip()


# ── Document generators ────────────────────────────────────────────

def _save_docx(text: str, path: str):
    doc = Document()
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)
    doc.save(path)


def _save_pdf(text: str, path: str):
    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleCustom", parent=styles["Title"], fontSize=18, spaceAfter=14)
    heading_style = ParagraphStyle("HeadingCustom", parent=styles["Heading2"], fontSize=13, spaceAfter=8)
    body_style = ParagraphStyle("BodyCustom", parent=styles["BodyText"], fontSize=11, leading=15)

    story = []
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 8))
        elif stripped.startswith("# "):
            story.append(Paragraph(stripped[2:], title_style))
        elif stripped.startswith("## ") or stripped.startswith("### "):
            clean = stripped.lstrip("#").strip()
            story.append(Paragraph(clean, heading_style))
        elif stripped.startswith("- "):
            story.append(Paragraph(f"• {stripped[2:]}", body_style))
        else:
            # Escape XML-sensitive chars
            safe = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            # Strip bold markdown markers
            safe = safe.replace("**", "")
            story.append(Paragraph(safe, body_style))
    doc.build(story)


# ── Public API ──────────────────────────────────────────────────────

def handle_proposal(client: OpenAI, session: dict, user_message: str, model="gpt-4o-mini"):
    """Process one turn of the proposal conversation.

    Returns (reply_text, proposal_id_or_None, is_complete).
    """
    conv = session.setdefault("proposal_conv", [])
    conv.append({"role": "user", "content": user_message})

    fields = _extract_fields(client, conv, model)
    # Merge with previously stored fields
    stored = session.get("proposal_fields", {})
    stored.update({k: v for k, v in fields.items() if v})
    session["proposal_fields"] = stored

    missing = _missing(stored)
    if missing:
        question = _ask_next(missing)
        conv.append({"role": "assistant", "content": question})
        return question, None, False

    # All required fields collected — generate proposal
    lang = session.get("language", "English")
    proposal_text = _generate_proposal_text(client, stored, model, lang)

    # Save documents
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    base = f"proposal_{stored['client_name'].replace(' ','_')}_{ts}"
    pdf_path = os.path.join(OUTPUT_DIR, f"{base}.pdf")
    docx_path = os.path.join(OUTPUT_DIR, f"{base}.docx")
    _save_pdf(proposal_text, pdf_path)
    _save_docx(proposal_text, docx_path)

    # Persist to database
    cid = db.get_or_create_client(stored["client_name"])
    pid = db.save_proposal(
        client_id=cid,
        client_name=stored["client_name"],
        project_title=stored["project_title"],
        proposal_text=proposal_text,
        deliverables=stored.get("deliverables", ""),
        timeline=stored.get("timeline", ""),
        budget=stored.get("budget", ""),
        freelancer_name=stored.get("freelancer_name", ""),
        freelancer_skills=stored.get("freelancer_skills", ""),
        file_path_pdf=pdf_path,
        file_path_docx=docx_path,
    )

    reply = (
        f"✅ **Proposal Generated!**\n\n{proposal_text}\n\n"
        f"---\n📥 Use the buttons below to download your proposal."
    )
    # Clear session state for proposals
    session.pop("proposal_conv", None)
    session.pop("proposal_fields", None)
    return reply, pid, True
